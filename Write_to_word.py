'''
CREATE WORD DOCUMENT WIDGET LAUNCHER developed by Mr Steven J walden
    Nov. 2019
    SAMROIYOD, PRACHUAP KIRI KHAN, THAILAND
[See license at end of file]

These classes are part of the main Write to word application

'''

__author__ = 'Steven J Walden'
__version__ = '1.0'

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from contextlib import contextmanager

@contextmanager
def change_dir(destination): #change directory function
	try:
		cwd = os.getcwd()
		os.chdir(destination)
		yield
	finally:
		os.chdir(cwd)

class Create_word_doc(object):
	"""docstring for Create_word_doc"""
	def __init__(self, table_list, file_name):
		super(Create_word_doc, self).__init__()

		#self.mylist = table_list
		self.mylist = table_list
		self.save_file_name = file_name

		self.mydoc = Document()
		self.header_table()
		self.heading('Your title goes here')
		self.p_graph('You can add text here in any format with Bullets or numbers etc')
		self.add_table()

		self.mydoc.save(self.save_file_name)

	def header_table(self):
		self.header_tbl = self.mydoc.add_table(rows=0, cols=2)
		self.header_cells = self.header_tbl.add_row().cells
		self.header_paragraph = self.header_cells[0].paragraphs[0]
		self.header_run = self.header_paragraph.add_run()
		self.header_run.bold = True
		self.header_run.add_text('Number address goes here')
		self.header_run.add_break()
		self.header_run.add_text('Street address goes here')
		self.header_run.add_break()
		self.header_run.add_text('Town address goes here')
		self.header_run.add_break()
		self.header_run.add_text('Postcode address goes here')

		self.paragraph = self.header_cells[1].paragraphs[0]
		self.run = self.paragraph.add_run()
		self.run_format = self.paragraph.paragraph_format
		self.run_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
		with change_dir('img'):
			self.run.add_picture('scared_robot.png', width=Mm(25.00))

	def heading(self, title):
		self.heading = self.mydoc.add_heading(title, 0)
		self.heading_format = self.heading.paragraph_format
		self.heading_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

	def p_graph(self, text):
		self.mydoc.add_paragraph(text, style=None)

	def add_table(self):
		self.table = self.mydoc.add_table(rows=1, cols=3)
		self.table_cells = self.table.rows[0].cells
		self.table_cells[0].text = 'Number'
		self.table_cells[1].text = 'Name'
		self.table_cells[2].text = 'Nickname'
		#read from a list
		for nm, nn, pw in self.mylist:
		    self.row_cells = self.table.add_row().cells
		    self.row_cells[0].text = nm
		    self.row_cells[1].text = nn
		    self.row_cells[2].text = pw


# Copyright (c) 2019 Steven Walden
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in
# all copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.